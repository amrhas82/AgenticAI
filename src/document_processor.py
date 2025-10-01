"""
Enhanced document processor supporting multiple file formats
Supports: PDF, TXT, MD, DOC, DOCX
"""

import io
from typing import List, Optional
import os

try:
    import PyPDF2
    _HAVE_PYPDF = True
except ImportError:
    _HAVE_PYPDF = False

try:
    from docx import Document as DocxDocument
    _HAVE_DOCX = True
except ImportError:
    _HAVE_DOCX = False


class DocumentProcessor:
    """Unified document processor for multiple file formats"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def process_file(self, file, filename: Optional[str] = None) -> List[str]:
        """
        Process a file and return text chunks
        
        Args:
            file: File object or path
            filename: Original filename (to determine file type)
        
        Returns:
            List of text chunks
        """
        if filename is None:
            filename = getattr(file, 'name', '')
        
        # Determine file type
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        
        try:
            if ext == 'pdf':
                return self.process_pdf(file)
            elif ext == 'docx':
                return self.process_docx(file)
            elif ext in ['txt', 'md', 'markdown', 'text']:
                return self.process_text(file)
            elif ext == 'doc':
                # .doc files require special handling (python-docx doesn't support them)
                # For now, return error message
                raise ValueError(
                    ".doc files are not supported. Please convert to .docx format. "
                    "You can use LibreOffice or MS Word to save as .docx"
                )
            else:
                # Try to process as plain text
                return self.process_text(file)
                
        except Exception as e:
            print(f"Error processing file {filename}: {e}")
            return []
    
    def process_pdf(self, pdf_file) -> List[str]:
        """Extract text from PDF and split into chunks"""
        if not _HAVE_PYPDF:
            raise ImportError("PyPDF2 is required for PDF processing. Install with: pip install pypdf2")
        
        try:
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            # Split into chunks
            chunks = self._split_text(text)
            return chunks
            
        except Exception as e:
            print(f"Error processing PDF: {e}")
            return []
    
    def process_docx(self, docx_file) -> List[str]:
        """Extract text from DOCX and split into chunks"""
        if not _HAVE_DOCX:
            raise ImportError(
                "python-docx is required for DOCX processing. "
                "Install with: pip install python-docx"
            )
        
        try:
            # Read DOCX
            document = DocxDocument(docx_file)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in document.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            # Split into chunks
            chunks = self._split_text(text)
            return chunks
            
        except Exception as e:
            print(f"Error processing DOCX: {e}")
            return []
    
    def process_text(self, text_file) -> List[str]:
        """Process plain text file and split into chunks"""
        try:
            # Handle both file objects and strings
            if isinstance(text_file, str):
                text = text_file
            elif hasattr(text_file, 'read'):
                content = text_file.read()
                # Decode if bytes
                if isinstance(content, bytes):
                    text = content.decode('utf-8', errors='ignore')
                else:
                    text = content
            else:
                raise ValueError("Unsupported text file format")
            
            # Split into chunks
            chunks = self._split_text(text)
            return chunks
            
        except Exception as e:
            print(f"Error processing text: {e}")
            return []
    
    def _split_text(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks by words
        
        Args:
            text: Text to split
        
        Returns:
            List of text chunks
        """
        if not text or not text.strip():
            return []
        
        # Split by words
        words = text.split()
        
        if len(words) <= self.chunk_size:
            return [text]
        
        chunks = []
        step = max(1, self.chunk_size - self.chunk_overlap)
        
        for i in range(0, len(words), step):
            chunk_words = words[i:i + self.chunk_size]
            if chunk_words:
                chunks.append(" ".join(chunk_words))
            
            # Stop if we've reached the end
            if i + self.chunk_size >= len(words):
                break
        
        return chunks
    
    def get_file_info(self, file, filename: Optional[str] = None) -> dict:
        """
        Get information about a file without processing it
        
        Returns:
            Dictionary with file metadata
        """
        if filename is None:
            filename = getattr(file, 'name', 'unknown')
        
        ext = filename.lower().split('.')[-1] if '.' in filename else 'unknown'
        
        # Try to get file size
        size = 0
        if hasattr(file, 'size'):
            size = file.size
        elif hasattr(file, 'seek') and hasattr(file, 'tell'):
            current_pos = file.tell()
            file.seek(0, 2)  # Seek to end
            size = file.tell()
            file.seek(current_pos)  # Restore position
        
        return {
            'filename': filename,
            'extension': ext,
            'size_bytes': size,
            'size_mb': round(size / (1024 * 1024), 2) if size > 0 else 0,
            'supported': ext in ['pdf', 'txt', 'md', 'markdown', 'docx', 'text']
        }


# Backward compatibility: keep PDFProcessor for existing code
class PDFProcessor(DocumentProcessor):
    """Legacy PDF processor - redirects to DocumentProcessor"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        super().__init__(chunk_size, chunk_overlap)
        print("PDFProcessor is deprecated. Use DocumentProcessor instead.")
    
    def process_pdf(self, pdf_file) -> List[str]:
        """Process PDF file - backward compatible method"""
        return super().process_pdf(pdf_file)
